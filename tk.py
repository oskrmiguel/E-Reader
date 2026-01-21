# -- coding: utf-8 --

import os
from PIL import Image, ImageTk 
import tkinter as tk
import cv2
from PIL import Image, ImageTk
import google.generativeai as genai  #GEMINI-1 
from gtts import gTTS  #Texto a audio
from langdetect import detect
from rich.console import Console #HTMLtexto consola
from rich.syntax import Syntax
import time


from io import BytesIO #almacenar img, audio 
import keyboard
import pygame
import pygame.camera
import re
import datetime


from docx import Document
import html2text 
import uuid

import subprocess
import glob
import datetime
import shutil



# Definición de variables globales

#contador para la función self.info()
contador = 0    

#variable para la función self.block_keys() y self.unblock_keys()
is_blocked = False       

class CameraApp: 

        
    def to_plain_string(self,text):
        # Eliminar caracteres especiales excepto letras, números y tildes
        cleaned_text = re.sub(r'[^\w\sáéíóúÁÉÍÓÚ.,;:?!¡¿%]', '', text)
        return cleaned_text 

    def contiene_palabra(self, cadena, palabra):
        # Verificar si la cadena no es None antes de intentar la comparación
        if cadena is not None:
            cadena = cadena.lower()  # Convertir la cadena a minúsculas
            palabra = palabra.lower()  # Convertir la palabra a minúsculas
            return palabra in cadena
        else:
            return False   

    def cerrar_ventana(self, event):
        self.master.destroy()  

    def hablar(self,texto):
        # Generar el archivo de audio usando gTTS y BytesIO
        print(texto)
        with BytesIO() as f:
            tts = gTTS(text=texto, lang='es', slow= False)
            tts.write_to_fp(f)
            f.seek(0)
            
            # Reproducir el audio usando pygame mixer
            pygame.mixer.init()
            pygame.mixer.music.load(f)
            pygame.mixer.music.play()
            
            # Esperar hasta que termine de reproducirse el audio
            while pygame.mixer.music.get_busy():
                if keyboard.is_pressed('Enter'):
                        pygame.mixer.music.stop()
                        time.sleep(0.1)
                        break
                continue
 

    def desconexion_audio(self):
        # Inicializar pygame
        pygame.init()

        try:
            # Iniciar la reproducción del audio
            pygame.mixer.music.load("desconexion.mpeg")
            pygame.mixer.music.play()
            # Esperar hasta que el audio termine de reproducirse
            while pygame.mixer.music.get_busy():
                pygame.time.Clock().tick(10)
        except pygame.error as e:
            print("Error al reproducir el audio:", e)
            self.error()
        finally:
            # Detener pygame
            pygame.quit()
 
    def guardado_exitoso(self):
        # Inicializar pygame
        pygame.init()
        try:
            # Iniciar la reproducción del audio
            pygame.mixer.music.load("guardado_exitoso.mpeg")
            pygame.mixer.music.play()

            # Esperar hasta que el audio termine de reproducirse
            while pygame.mixer.music.get_busy():
                pygame.time.Clock().tick(10)
        except pygame.error as e:
            print("Error al reproducir el audio:", e)
            self.error()
        finally:
            # Detener pygame
            pygame.quit() 
 

    def error(self):
        pygame.init()
        try:
            pygame.mixer.music.load("error.mpeg")
            pygame.mixer.music.play()
            while pygame.mixer.music.get_busy():
                pygame.time.Clock().tick(10)
        except pygame.error as e:
            print("Error al reproducir el audio:", e)
            self.error()
        finally:
            pygame.quit() 


    def html_to_docx(self,text):
                # Carpeta donde se guardarán los documentos
                folder_path = "respuestas"
                if not os.path.exists(folder_path):
                    os.makedirs(folder_path)  
                output_path = os.path.join(folder_path, "unl-text"+str(uuid.uuid4())+".docx")  
		 
                # Crear un nuevo documento de Word
                doc = Document()
                texto_sin_saltos = text.replace("\n", " ")
                text_content = html2text.html2text(texto_sin_saltos) 
                doc.add_paragraph(text_content)
                doc.save(output_path)
                print("Guardado en "+folder_path+" como: unl-text"+str(uuid.uuid4())+".docx")

                try:
                    usb_path = os.path.join("/media/unl/KINGSTON", "unl-text"+str(uuid.uuid4())+".docx") 
                    doc = Document()
                    texto_sin_saltos = text.replace("\n", " ")
                    text_content = html2text.html2text(texto_sin_saltos) 
                    doc.add_paragraph(text_content)
                    doc.save(usb_path)
                except Exception as e:
                    print("No se pudo guardar el archivo en un USB nombrado como Kingston")
                    #self.error()

#https://gtts.readthedocs.io/en/latest/cli.html
    def reproducir_audio(self,texto):
        try:
            with BytesIO() as f:
                texto_sin_saltos = texto.replace("\n", " ") 
                #tts = gTTS(text=texto_sin_saltos, lang='pt')
                lang =detect(texto_sin_saltos)
                tts = gTTS(text=texto_sin_saltos, lang=lang)
                tts.write_to_fp(f)

                f.seek(0) 
                pygame.mixer.init()
                pygame.mixer.music.load(f)
                pygame.mixer.music.play()
                
                while pygame.mixer.music.get_busy():
                    if keyboard.is_pressed('0'):
                        pygame.mixer.music.stop()
                        time.sleep(0.1)
                        self.desconexion_audio()
                        break
                    
                    if keyboard.is_pressed('+'):
                        pos=pygame.mixer.music.get_pos()
                        pygame.mixer.music.set_pos(pos + 1) #adelantar un segundo 
                        
                    if keyboard.is_pressed('-'):
                        pos=pygame.mixer.music.get_pos()
                        pygame.mixer.music.set_pos(pos - 1) #adelantar un segundo 
                        
                    # Detectar si se ha pulsado la tecla "Enter"
                    pausado = False
                    if pausado == False:

                        #print("Esperando el . para pause")
                        if keyboard.is_pressed('Enter'):
                            pausado = True
                            pygame.mixer.music.pause()
                            pygame.mixer.music.pause()
                            #time.sleep(0.1)
                            print("pause")
                        elif keyboard.is_pressed('0'):
                            pygame.mixer.music.stop()
                            time.sleep(0.1) 
                            break 
                        while pausado == True:
                                print("Esperando el . para play")
                                if keyboard.is_pressed('Enter'):
                                    pausado==False 
                                    pygame.mixer.music.unpause()
                                    #time.sleep(0.1)
                                    print("Play") 
                                    break

                                elif keyboard.is_pressed('0'):
                                    pygame.mixer.music.stop()
                                    break
                    continue    
                self.preguntar_guardar(texto,tts) 
        except Exception as e:
            print("reproducir_audio(); GTTS ERROR - - ERROR - - ERROR   GTTS")
            self.error()
            print('{type(e).__name__}: '+ str(e))     

 


    def reproducir_audio_no_save(self,texto):
        try:
            with BytesIO() as f:
                texto_sin_saltos = texto.replace("\n", " ")  
                lang =detect(texto_sin_saltos)
                tts = gTTS(text=texto_sin_saltos, lang=lang)
                tts.write_to_fp(f)

                f.seek(0) 
                pygame.mixer.init()
                pygame.mixer.music.load(f)
                pygame.mixer.music.play()
                
                while pygame.mixer.music.get_busy():
                    if keyboard.is_pressed('0'):
                        pygame.mixer.music.stop()
                        time.sleep(0.1)
                        self.desconexion_audio()
                        break
                    
                    if keyboard.is_pressed('+'):

                        pos=pygame.mixer.music.get_pos()
                        pygame.mixer.music.set_pos(pos + 1) #adelantar un segundo 
                        
                    if keyboard.is_pressed('-'):
                        pos=pygame.mixer.music.get_pos()
                        pygame.mixer.music.set_pos(pos - 1) #adelantar un segundo 
                        
                    # Detectar si se ha pulsado la tecla "Enter"
                    pausado = False
                    if pausado == False:

                        #print("Esperando el . para pause")
                        if keyboard.is_pressed('Enter'):
                            pausado = True
                            pygame.mixer.music.pause()
                            #time.sleep(0.1)
                            print("pause")
                        elif keyboard.is_pressed('0'):
                            pygame.mixer.music.stop()
                            time.sleep(0.1) 
                            break 
                        while pausado == True:

                                print("Esperando el . para play")
                                if keyboard.is_pressed('Enter'):
                                    pausado==False 
                                    pygame.mixer.music.unpause()
                                    #time.sleep(0.1)
                                    print("Play") 
                                    break

                                elif keyboard.is_pressed('0'):
                                    pygame.mixer.music.stop()
                                    break
                    continue     
        except Exception as e:
            print("reproducir_audio(); GTTS ERROR - - ERROR - - ERROR   GTTS")
            self.error()
            print('{type(e).__name__}: '+ str(e))     


#https://gtts.readthedocs.io/en/latest/cli.html
    def reproducir_audio_nosound(self,texto):
        try:
            with BytesIO() as f:
                texto_sin_saltos = texto.replace("\n", " ") 
                lang =detect(texto_sin_saltos)
                tts = gTTS(text=texto_sin_saltos, lang=lang)
                tts.write_to_fp(f)

                f.seek(0) 
                pygame.mixer.init()
                pygame.mixer.music.load(f)
                pygame.mixer.music.play()
                
                while pygame.mixer.music.get_busy():
                    if keyboard.is_pressed('0'):
                        pygame.mixer.music.stop()
                        time.sleep(0.1) 
                        break
                    continue    

        except Exception as e:
            print("reproducir_audio(); GTTS ERROR - - ERROR - - ERROR   GTTS")
            self.error()
            print('{type(e).__name__}: '+ str(e))     

 

    def preguntar_guardar(self,texto,tts):
        try:
            self.reproducir_audio_nosound("Guardar texto tecla 7, audio 8, o ambos 9")
            tecla = keyboard.read_key()  
            if tecla=="7":
                self.html_to_docx(texto)    
                self.guardado_exitoso()
            if tecla=="8":
                tts = tts
                # Obtener la fecha y hora actual
                now = datetime.datetime.now()
                name = "respuestas/unl-audio-"+str(uuid.uuid4()) + ".mp3"

                # Carpeta donde se guardarán los audios 
                folder_path = "respuestas"
                if not os.path.exists(folder_path):
                    os.makedirs(folder_path)   
                tts.save(name)  
                
                try:
                    # USB donde se guardarán los audios   
                    nombre_audio_usb="/home/unl/unl2024/respuestas/audios/audio-"+str(uuid.uuid4())+".mp3"
                    tts.save(nombre_audio_usb)
                    print("Guardado en KINGSTON")
                except Exception as e:              
                    try:
                        # USB donde se guardarán los audios   
                        nombre_audio_usb="/home/unl/unl2024/respuestas/audios/audio-"+str(uuid.uuid4())+".mp3"
                        tts.save(nombre_audio_usb)
                        print("Guardado en kingston")
                    except Exception as e:
                        print("No se guardó en kingston o KINGSTON")
                self.guardado_exitoso()
            if tecla=="9":
                self.html_to_docx(texto) 

                tts = tts
                # Obtener la fecha y hora actual
                now = datetime.datetime.now()

                # Carpeta donde se guardarán los audios 
                folder_path = "respuestas"
                if not os.path.exists(folder_path):
                    os.makedirs(folder_path)  
                nombre_audio = "respuestas/unl-audio-"+str(uuid.uuid4()) + ".mp3"
                tts.save(nombre_audio)
                try:
                    # USB donde se guardarán los audios   
                    nombre_audio_usb="/media/unl/KINGSTON/unl-audio-"+str(uuid.uuid4())+".mp3"
                    tts.save(nombre_audio_usb)
                    print("Guardado en KINGSTON")
                except Exception as e:              
                    try:
                        # USB donde se guardarán los audios   
                        nombre_audio_usb="/media/unl/kingston/unl-audio-"+str(uuid.uuid4())+".mp3"
                        tts.save(nombre_audio_usb)
                        print("Guardado en kingston")
                    except Exception as e:

                        print("No se guardó en kingston o KINGSTON")

                self.guardado_exitoso()

            elif not tecla=="0":
                exit 
        
        except Exception as e:
            print("Error en la guardada del archivo")
            self.error()
            print('{type(e).__name__}: '+ str(e))     

 




    def gemini_unl(self, mensaje):
                foto, path =self.tomarFoto() 


                #Eliminar Foto del repositorio 
                os.remove(path)  
                print(" mensaje ",mensaje)  
                print(" foto ",foto)    
                model = genai.GenerativeModel('gemini-2.5-flash')  
                # Seguridad https://ai.google.dev/docs/safety_setting_gemini?hl=es-419

                safety_settings = {
                    "HARM_CATEGORY_HARASSMENT": "BLOCK_NONE",
                    "HARM_CATEGORY_HATE_SPEECH": "BLOCK_NONE",
                    "HARM_CATEGORY_SEXUALLY_EXPLICIT": "BLOCK_NONE",
                    "HARM_CATEGORY_DANGEROUS_CONTENT": "BLOCK_NONE"} 
                try:
                    response = model.generate_content([mensaje, foto], safety_settings=safety_settings)
                    #time.sleep(1)   
                    response.resolve() 
                    self.text_label.config(text=response)
                    print(response)
                    if response.candidates[0].content.parts: 
                        try: 
                            respuesta= response.candidates[0].content.parts[0].text   
                            respuesta= self.to_plain_string(respuesta) 
                            #Devolviendo respuesta
                            print(respuesta) 
                            self.text_label.config(text=respuesta)
                            return(respuesta)  
                        except Exception as e:
                            print("Entró al Primer CATCH, ERROR - - ERROR - - ERROR")
                            print('f{type(e).__name__}: ',{e})     
                            try:  # Si solo hay una respuesta: 
                                print("!!!!!!!!!!!!!!!")
                                a=response.text    
                                respuestab=self.to_plain_string(a)
                                #Devolviendo respuesta
                                self.text_label.config(text=respuestab)
                                return(respuestab)  

                            except Exception as e:
                                print("Entró al segundo CATCH, ERROR - - ERROR - - ERROR")
                                self.hablar("No se pudo generar una respuesta")
                                self.text_label.config(text="No se pudo generar una respuesta")
                                print('{type(e).__name__}: '+{e})     

                    else: # Si solo hay una respuesta:
                        try:  
                            print("Parece que solo hay una respuesta. :)")
                            print("!!!!!!!!!!!!!!!")
                            a=response.text   
                            print("response.text")
                            respuesta=self.to_plain_string(respuesta)  
                            #Devolviendo respuesta
                            self.text_label.config(text=respuesta)
                            return(respuesta)  

                        except Exception as e:
                            print("Entró al segundo CATCH, del Else")
                            print('{type(e).__name__}: '+{e})     
                            try:  # Si solo hay una respuesta:
                                print("Parece que solo hay una respuesta. :)")  
                                respuesta= response.candidates[0].content.parts[0].text 
                                #Devolviendo respuesta
                                return(respuesta)  

                            except Exception as e:
                                print("ERROR - - ERROR")
                                self.error()
                                print('f{type(e).__name__}: ',+str(e))     
                    
                except Exception as e:
                    print("No se pudo generar una respuesta")
                    self.error() 
                    print('f{type(e).__name__}: ',{e})    
                

#####
    def gemini(self, prompt):
            self.start_function()  
            prompt=prompt
            if self.contiene_palabra(prompt, "Identifica"):   
                print("Ejecutar sin guardar- (Identificación del contenido del documento)")
                try:  
                    res = self.gemini_unl(prompt)  

                    #REPRODUCIR 
                    self.reproducir_audio_no_save(res)  
                    self.desconexion_audio()
                    self.executing = False
                    self.finish_function()  

                except Exception as e:
                    print("Hubo un error en la ultima funcion!")
                    self.error()
                    self.executing = False
                    self.finish_function()  
                    print('{type(e).__name__}: '+str(e))    
            else: 
                print("Ejecutar y guardar - (Lectura del documento)") 
                try:  
                    res = self.gemini_unl(prompt) 
                    #REPRODUCIR 
                    self.reproducir_audio(res)  
                    self.desconexion_audio()
                    self.executing = False
                    self.finish_function()  

                except Exception as e:
                    print("Hubo un error en la ultima funcion!")
                    self.error()
                    self.executing = False
                    self.finish_function()  
                    print('{type(e).__name__}:'+str(e))  
########
#Ver logs entiempo reaL ##     dmesg -w
#Ver logs del sistema   ##     journalctl -k

    def __init__(self, window, window_title):  

            GOOGLE_API_KEY='' #Aquí la clave del api key para concetarme con gemini
            genai.configure(api_key=GOOGLE_API_KEY) 
            self.window = window
            self.window.title(window_title)
            self.window.attributes('-fullscreen', True)

            #Superponer la ventana
            #self.window.attributes('-topmost',True)    
 
            # Obtener el tamaño de la pantalla
            #screen_width = self.window.winfo_screenwidth()
            #screen_height = self.window.winfo_screenheight()

            # Ajustar el tamaño de la ventana a la pantalla
            #self.window.geometry(f"{screen_width}x{screen_height}")

            #Cámara por defecto de la PC

            #self.vid = cv2.VideoCapture(0) 
            #self.canvas = tk.Canvas(window, width=self.vid.get(cv2.CAP_PROP_FRAME_WIDTH), height=self.vid.get(cv2.CAP_PROP_FRAME_HEIGHT))
            #self.canvas.pack()  
            #self.canvas.place(x=0, y=60) 
 

            # Cargar y ajustar la imagen de fondo
            #self.update_background_image()

        # Crear un widget Label para mostrar la imagen
            self.image_label = tk.Label(window)
            self.image_label.pack(fill=tk.BOTH, expand=True)

                # IMAGEN BACKGROUND        
            self.original_image = Image.open("background.jpeg")

        # Redimensionar la imagen según el tamaño inicial de la ventana
            self.update_image()
# Asociar el método de actualización de la imagen al evento de cambio de tamaño
            self.window.bind("<Configure>", self.on_resize)
 
        # Crear el Label con el texto
            #inicialx=self.window.winfo_screenwidth()*0.47
            #inicialy=self.window.winfo_screenheight()*0.213
            #anchotexto=self.window.winfo_screenwidth()/3.36
            inicialx=self.window.winfo_screenwidth()*0.6
            inicialy=self.window.winfo_screenheight()*0.23
            anchotexto=self.window.winfo_screenwidth()/2.64
            self.text_label = tk.Label(self.image_label, text="Bienvenido a E-Reader", fg="white", bg="black", font=("Arial", 9), wraplength=anchotexto, justify="left")
            self.text_label.place(x=inicialx, y=inicialy)  # Cambia las coordenadas según necesites
            self.add_buttons()
            #self.update()
 
  

## COMANDS  >>> 
            self.executing = False
            self.window.bind('<Escape>', lambda e: self.exit()) 
            self.window.bind('<KeyPress>', self.on_key_press) 
            self.window.mainloop()

    def add_buttons(self):
        inicial=self.window.winfo_screenwidth()-180
        # Botón *
        self.buttonn = tk.Button(self.image_label, text="*", command=self.button_actionn, bg='black', fg='white')
        self.buttonn.place(x=inicial-330, y=115) 
        
        # Botón /
        self.butto = tk.Button(self.image_label, text="/", command=self.button_actio, bg='black', fg='white')
        self.butto.place(x=inicial-300, y=115)  
        # Botón 1
        self.button1 = tk.Button(self.image_label, text="1", command=self.button1_action, bg='black', fg='white')
        self.button1.place(x=inicial-270, y=115)
        
        # Botón 2
        self.button2 = tk.Button(self.image_label, text="2", command=self.button2_action, bg='black', fg='white')
        self.button2.place(x=inicial-240, y=115) 
        # Botón 3
        self.button3 = tk.Button(self.image_label, text="3", command=self.button3_action, bg='black', fg='white')
        self.button3.place(x=inicial-210, y=115)
        
        # Botón 4
        self.button4 = tk.Button(self.image_label, text="4", command=self.button4_action, bg='black', fg='white')
        self.button4.place(x=inicial-180, y=115) 

        # Botón 5
        self.button5 = tk.Button(self.image_label, text="5", command=self.button5_action, bg='black', fg='white')
        self.button5.place(x=inicial-150, y=115) 
        
        # Botón 7
        self.button7 = tk.Button(self.image_label, text="7", command=self.button7_action, bg='white', fg='black')
        self.button7.place(x=inicial-120, y=115)
        # Botón 8
        self.button8 = tk.Button(self.image_label, text="8", command=self.button8_action, bg='white', fg='black')
        self.button8.place(x=inicial-90, y=115) 
        
        # Botón 9
        self.button9 = tk.Button(self.image_label, text="9", command=self.button9_action, bg='white', fg='black')
        self.button9.place(x=inicial-60, y=115)
        # Botón 0
        self.button0 = tk.Button(self.image_label, text="0", command=self.button0_action, bg='gray', fg='white')
        self.button0.place(x=inicial, y=115)
        
        # Botón Enter
        self.buttonE = tk.Button(self.image_label, text="Enter", command=self.buttonE_action, bg='gray', fg='white')
        self.buttonE.place(x=inicial+30, y=115)  
 
    def button_actionn(self):
        print("Botón * presionado")
        self.start_function()
        print("Tecla '*' presionada")
        is_blocked = True       
        self.incrementar_contador()
        self.info                                                                ()
        self.executing = False
        
    def button_actio(self):
        print("Botón / presionado") 
        self.start_function()
        print("Tecla '/' presionada")
        is_blocked = True      
        self.decrementar_contador()
        self.info()
        self.executing = False 
 

    def button1_action(self):
        print("Botón 1 presionado")
        self.button1.focus_set()
        self.gemini("Identifica que hay en la imagen, hay texto?, tablas? graficos? en español: solo responde algo como: el contenido del documento es texto, el contenido del documento es gráficos, o el contenido del documento es tablas, o el contenido del documento es texto y tablas o texto y gráficos, o si hay texto tablas y gráficos, o no hay, nada, dependiendo del contenido del documento, no des una explicación del documento.")

    def button2_action(self):
        print("Botón 2 presionado") 
        self.gemini("extrae el texto que encuentres")
           
    def button3_action(self):
        print("Botón 3 presionado")
        self.gemini("Narra en español todo el contenido de la tabla de una forma ordenada, coherente y estructurada de forma que se pueda entender TODO el contenido, primero narra las columnas y luego narra los valores de cada fila relacionando con cada columna")

    def button4_action(self):
        print("Botón 4 presionado") 
        self.gemini("Describe en español todo el contenido del gráfico o gráficos estadíscos presentes en la imagen, y leelo de una forma ordenada, coherente y estructurada de forma que se pueda entender TODO el contenido")
    def button5_action(self):
        print("Botón 5 presionado")     
        self.gemini("Extrae el texto de los párrafos y Describe todo el contenido, si hay tablas o gráfico, describe cada fila o sección explicando de forma ordenada, coherente y estructurada de forma que se pueda entender TODO el contenido")

    def button7_action(self):
        print("Botón 7 presionado") 
            
    def button8_action(self):
        print("Botón 8 presionado")

    def button9_action(self):
        print("Botón 9 presionado") 

    def button0_action(self):
        print("Botón 0 presionado")

    def buttonE_action(self):
        print("Botón Enter presionado") 

            
    def capturarFoto(self):
        ret, frame = self.vid.read()
        if ret:
            retval, buffer = cv2.imencode('.jpg', frame)
            img_bytes = BytesIO(buffer.tobytes())
            img_pil = Image.open(img_bytes) 
            return img_pil
        else:
            print("Error al capturar la foto.")

    #def update(self):
     #   ret, frame = self.vid.read()
      #  if ret:
       #     self.photo = ImageTk.PhotoImage(image=Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)))
        #    self.canvas.create_image(0, 0, image=self.photo, anchor=tk.NW)
         #   self.window.after(self.delay, self.update)


    def on_key_press(self, event): 
        global is_blocked
        # Verifica si el teclado está bloqueado
        if is_blocked:
            return  # No hace nada si el teclado está bloqueado

        # Manejo de eventos de teclas
        if event.keysym in ['0', 'KP_0']:
            print("Tecla '0' presionada")
            self.executing = False
        elif event.keysym in ['1', 'KP_1']:
            print("Tecla '1' presionada")
            self.button1_action() 
            self.executing = False
        elif event.keysym in ['2', 'KP_2']:
            print("Tecla '2' presionada")
            self.button2_action() 
            self.executing = False
        elif event.keysym in ['3', 'KP_3']:
            print("Tecla '3' presionada") 
            self.button3_action() 
            self.executing = False
        elif event.keysym in ['4', 'KP_4']:
            print("Tecla '4' presionada")
            self.button4_action() 
            self.executing = False
        elif event.keysym in ['5', 'KP_5']:
            print("Tecla '5' presionada")
            self.button5_action() 
            self.executing = False
        elif event.keysym in ['6', 'KP_6']:
            print("Tecla '6' presionada")
            self.gemini("Describe detalladamente la imagen")
            self.executing = False
        elif event.keysym in ['7', 'KP_7']:
            print("Tecla '7' presionada")
            self.button7_action() 
            self.executing = False
        elif event.keysym in ['8', 'KP_8']:
            print("Tecla '8' presionada")
            self.executing = False
        elif event.keysym in ['9', 'KP_9']:
            print("Tecla '9' presionada")
            self.executing = False
        elif event.keysym in ['asterisk', 'KP_Multiply']:
            self.start_function()
            print("Tecla '*' presionada")
            is_blocked = True       
            self.incrementar_contador()
            self.info()
            self.executing = False
        elif event.keysym in ['slash', 'KP_Divide']:
            self.start_function()
            print("Tecla '/' presionada")
            is_blocked = True      
            self.decrementar_contador()
            self.info()
            self.executing = False 
 

    def update_image(self):
        # Obtener el tamaño de la ventana
        width = self.window.winfo_width()
        height = self.window.winfo_height()
        
        # Redimensionar la imagen
        resized_image = self.original_image.resize((width, height), Image.LANCZOS)
        
        # Convertir la imagen redimensionada a un objeto PhotoImage
        self.photo = ImageTk.PhotoImage(resized_image)
        
        # Actualizar la imagen en el widget Label
        self.image_label.config(image=self.photo)
        self.image_label.image = self.photo  # Mantener una referencia para evitar que la imagen sea recogida por el garbage collector

    def on_resize(self, event):
        # Llamar al método de actualización de la imagen cuando cambie el tamaño de la ventana
        self.update_image()

            
    #def capturarFoto(self):
    #    ret, frame = self.vid.read()
    #    if ret:
    #        retval, buffer = cv2.imencode('.jpg', frame)
    #        img_bytes = BytesIO(buffer.tobytes())
    #        img_pil = Image.open(img_bytes) 
    #        return img_pil
    #    else:
    #        print("Error al capturar la foto.")

 #   def update(self):
 #       ret, frame = self.vid.read()
 #       if ret:
 #           self.photo = ImageTk.PhotoImage(image=Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)))
 #           self.canvas.create_image(0, 0, image=self.photo, anchor=tk.NW)
 #           self.window.after(self.delay, self.update)

 

## COMANDS <<< 



    def info(self):
        global contador
        if contador == 1:
            self.reproducir_audio_no_save("Bienvenido a Eye Reader")
            self.finish_function()
        elif contador == 2:
            self.reproducir_audio_no_save("Para iniciar, coloca un documento en la bandeja")
            self.finish_function()
        elif contador == 3:
            self.reproducir_audio_no_save("Con la tecla 1 identifica el tipo de contenido")
            self.finish_function()
        elif contador == 4:
            self.reproducir_audio_no_save("La tecla 2 ejecuta la lectura de texto")
            self.finish_function()
        elif contador == 5:
            self.reproducir_audio_no_save("La tecla 3 describe las tablas presentes en el documento")
            self.finish_function()
        elif contador == 6:
            self.reproducir_audio_no_save("La tecla 4 describe los gráficos")
            self.finish_function()
        elif contador == 7:
            self.reproducir_audio_no_save("La tecla 5 para leer texto con tablas y/o gráficos")
            self.finish_function()
        elif contador == 8:
            self.reproducir_audio_no_save("Puedes guardar el texto o audio generado; Despues de cada ejecución presiona 7 para guardar el texto, 8 para guardar el audio o 9 para guardar ambos")
            self.resetear_contador()
            self.finish_function()   
        else:
            self.resetear_contador()

    def start_function(self):
        global is_blocked
        # Bloquear el teclado
        self.is_blocked = True
        is_blocked = True
        print("Teclado bloqueado")

    def finish_function(self):
        global is_blocked
        # Termina la función y desbloquea el teclado
        self.is_blocked = False 
        is_blocked = False
        print("Teclado desbloqueado")

    def incrementar_contador(self):
        # Indicar que se usará la variable global 'contador'
        global contador
        contador += 1
        print(f'El contador es ahora: {contador}')

    def decrementar_contador(self):
        # Indicar que se usará la variable global 'contador'
        global contador
        contador = contador-1
        print(f'El contador es ahora: {contador}')

    def resetear_contador(self):
        global contador
        contador = 0
        print('El contador ha sido reiniciado')
 


    def tomarFoto(self):
        self.subprocess()
        # Busca todos los archivos JPG en el directorio
        archivos_jpg = glob.glob(os.path.join(".", "*.jpg"))

        # Ordena los archivos por fecha de creación
        archivos_jpg.sort(key=os.path.getmtime)
        
        # Obtiene el último archivo creado
        ultimo_archivo_jpg = archivos_jpg[-1]
        
        # Abre el último archivo JPG
        imagen = Image.open(ultimo_archivo_jpg)  

	#Colocar la imagen en el fondo 
        #overlay_photo = ImageTk.PhotoImage(imagen) 
        #self.text_label.config(image=overlay_photo)
        #self.text_label = tk.Label(self.image_label, image=overlay_photo)
        #self.text_label.place(x=50, y=50)  # Ajusta la posición según sea necesario
        #os.remove(ultimo_archivo_jpg)
         
        return (imagen,ultimo_archivo_jpg) 

    def subprocess(self): 
        #subprocess.run(["nvgstcapture-1.0", "--automate", "--capture-auto", "--image-res", "3"], shell=False, capture_output=True, text=True, check=True)
        subprocess.run(["nvgstcapture-1.0", "--automate", "--capture-auto", "--image-res", "3"], shell=False, capture_output=True, text=True, check=True)
#Supported resolutions in case of NvArgusCamera cambiar en --image-res
#  (2) : 640x480  MAL
#  (3) : 1280x720 GOD
#  (4) : 1920x1080  (abarca toda la hoja sin bordes)
#  (5) : 2104x1560
#  (6) : 2592x1944
#  (7) : 2616x1472
#  (8) : 3840x2160  (abarca toda la hoja sin bordes)
#  (9) : 3896x2192
#  (10): 4208x3120
#  (11): 5632x3168
#  (12): 5632x4224


    def exit(self): 
        self.window.destroy()

App = CameraApp(tk.Tk(), "Lector de documentos - UNL")
